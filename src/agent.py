"""
================================================================================
TALENTMATCH AI — HR Resume Shortlisting Agent
================================================================================
Stack  : Gemini 2.5 Flash · LangChain LCEL · TF-IDF · LIME XAI · ReportLab
Author : [Your Name] | 3rd Year B.Tech AIML
Task   : AI Enablement Internship – Task 1
================================================================================
"""

# ── Stdlib ────────────────────────────────────────────────────────────────────
import os, re, json, logging, time
from datetime import datetime
from typing import List, Dict, Optional

# ── Document parsing ──────────────────────────────────────────────────────────
import pypdf
from docx import Document as DocxDocument

# ── Data / ML ─────────────────────────────────────────────────────────────────
import pandas as pd
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# ── XAI ───────────────────────────────────────────────────────────────────────
import lime
import lime.lime_text

# ── LangChain LCEL ────────────────────────────────────────────────────────────
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.globals import set_llm_cache
from langchain_community.cache import SQLiteCache

# ── PDF report ────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle,
    Paragraph, Spacer, HRFlowable, PageBreak,
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────
DELAY_BETWEEN_CANDIDATES = 4   # seconds — stays within Gemini free-tier RPM

# Brand palette
C_NAVY   = colors.HexColor("#1B3A6B")
C_TEAL   = colors.HexColor("#0D9488")
C_AMBER  = colors.HexColor("#F59E0B")
C_PURPLE = colors.HexColor("#7C3AED")
C_LGRAY  = colors.HexColor("#F3F4F6")
C_DGRAY  = colors.HexColor("#374151")

# ─────────────────────────────────────────────────────────────────────────────
# AUDIT LOGGER  (Security: full trace of every agent action)
# ─────────────────────────────────────────────────────────────────────────────
logging.basicConfig(
    filename="audit_log.txt",
    level=logging.INFO,
    format="%(asctime)s | %(message)s",
)

def audit(event: str, detail: str = "") -> None:
    """Append a structured entry to audit_log.json + audit_log.txt."""
    entry = {"timestamp": datetime.now().isoformat(), "event": event, "detail": detail}
    logging.info(f"{event} | {detail}")
    with open("audit_log.json", "a") as f:
        f.write(json.dumps(entry) + "\n")


# ─────────────────────────────────────────────────────────────────────────────
# PII MASKING  (Security: never write raw emails/phones to logs)
# ─────────────────────────────────────────────────────────────────────────────
def mask_pii(text: str) -> str:
    """Mask emails and phone numbers before writing to any log."""
    # email: first char + **** + domain
    text = re.sub(
        r"([\w.+-])([\w.+-]+)(@[\w-]+\.[a-zA-Z]+)",
        lambda m: m.group(1) + "****" + m.group(3),
        text,
    )
    text = re.sub(r"\b(\+?\d[\d\s\-]{7,}\d)\b", "[PHONE]", text)
    return text


# ─────────────────────────────────────────────────────────────────────────────
# LLM SETUP  (LangChain + SQLite caching)
# ─────────────────────────────────────────────────────────────────────────────
def init_llm(api_key: str) -> ChatGoogleGenerativeAI:
    """Initialise Gemini via LangChain with SQLite response caching."""
    os.environ["GOOGLE_API_KEY"] = api_key
    set_llm_cache(SQLiteCache(database_path="llm_cache.db"))
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash-preview-05-20",
        google_api_key=api_key,
        temperature=0,
        max_tokens=2048,
        max_retries=3,
    )
    audit("LLM_INIT", "model=gemini-2.5-flash | cache=SQLite")
    return llm


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT READERS
# ─────────────────────────────────────────────────────────────────────────────
def read_pdf(path: str, max_mb: int = 10) -> str:
    assert os.path.getsize(path) / 1e6 <= max_mb, f"PDF > {max_mb} MB"
    text = ""
    with open(path, "rb") as f:
        for page in pypdf.PdfReader(f).pages:
            text += page.extract_text() or ""
    return text.strip()


def read_docx(path: str, max_mb: int = 10) -> str:
    assert os.path.getsize(path) / 1e6 <= max_mb, f"DOCX > {max_mb} MB"
    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(c.text for c in row.cells))
    return "\n".join(paragraphs)


def read_resume(path: str) -> str:
    """Universal reader: PDF / DOCX / TXT."""
    ext = path.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return read_pdf(path)
    elif ext in ("docx", "doc"):
        return read_docx(path)
    elif ext == "txt":
        with open(path, encoding="utf-8") as f:
            return f.read()
    raise ValueError(f"Unsupported format: .{ext}")


def read_linkedin_json(path: str) -> str:
    """Convert exported LinkedIn JSON to plain text for the pipeline."""
    with open(path) as f:
        d = json.load(f)
    name  = f"{d.get('firstName','')} {d.get('lastName','')}".strip()
    skills = ", ".join(d.get("skills", []))
    certs  = ", ".join(d.get("certifications", []))
    summary = d.get("summary", "")
    exp_text = ""
    for p in d.get("positions", []):
        s = p.get("startDate", {}).get("year", "")
        e = p.get("endDate", {}).get("year", "Present") if p.get("endDate") else "Present"
        exp_text += f"\n- {p.get('title')} at {p.get('companyName')} ({s}–{e}): {p.get('description','')}"
    edu_text = ""
    for e in d.get("educations", []):
        edu_text += f"\n- {e.get('degreeName','')} in {e.get('fieldOfStudy','')} from {e.get('schoolName','')}"
    return (
        f"Name: {name}\nSummary: {summary}\nSkills: {skills}\n"
        f"Certifications: {certs}\nExperience:{exp_text}\nEducation:{edu_text}\n"
        f"Source: LinkedIn Profile"
    )


# ─────────────────────────────────────────────────────────────────────────────
# JD PARSER  (LangChain LCEL)
# ─────────────────────────────────────────────────────────────────────────────
_JD_PROMPT = PromptTemplate(
    input_variables=["jd_text"],
    template="""You are an expert HR analyst. Extract structured requirements.

Job Description:
{jd_text}

Return ONLY valid JSON — no markdown, no explanation:
{{
  "role_title": "exact job title",
  "required_skills": ["skill1", "skill2"],
  "nice_to_have": ["skill1"],
  "min_experience": 2,
  "max_experience": 5,
  "required_education": "B.Tech/M.Tech in CS or related",
  "domain": "fintech",
  "seniority": "Senior"
}}

JSON:""",
)

def analyze_jd(jd_text: str, llm) -> dict:
    """Parse JD → structured requirements dict. Cached in SQLite."""
    chain = _JD_PROMPT | llm | StrOutputParser()
    raw = chain.invoke({"jd_text": jd_text})
    raw = raw.strip().replace("```json", "").replace("```", "").strip()
    try:
        result = json.loads(raw)
        audit("JD_PARSED", f"role={result.get('role_title')} skills={len(result.get('required_skills',[]))}")
        return result
    except json.JSONDecodeError:
        audit("JD_PARSE_ERROR", raw[:120])
        return {"role_title": "Unknown", "required_skills": [], "nice_to_have": [],
                "min_experience": 0, "domain": "General", "seniority": "Mid",
                "required_education": "Bachelor's Degree"}


# ─────────────────────────────────────────────────────────────────────────────
# CANDIDATE PROFILER  (LangChain LCEL)
# ─────────────────────────────────────────────────────────────────────────────
_PROFILE_PROMPT = PromptTemplate(
    input_variables=["resume_text"],
    template="""Extract candidate information from this resume.

Resume:
{resume_text}

Return ONLY valid JSON:
{{
  "name": "Full Name",
  "email": "email or null",
  "skills": ["skill1", "skill2"],
  "total_experience": 3.5,
  "education": ["degree, institution, year"],
  "certifications": ["cert1"],
  "work_history": [{{"company": "", "role": "", "duration": "", "domain": ""}}],
  "projects": ["project description"],
  "summary": "one-line summary"
}}

JSON:""",
)

def extract_candidate(resume_text: str, llm) -> dict:
    chain = _PROFILE_PROMPT | llm | StrOutputParser()
    raw = chain.invoke({"resume_text": resume_text[:4000]})
    raw = raw.strip().replace("```json", "").replace("```", "").strip()
    try:
        return json.loads(raw)
    except Exception:
        return {"name": "Unknown", "email": None, "skills": [],
                "total_experience": 0, "education": [], "certifications": [],
                "work_history": [], "projects": [], "summary": ""}


# ─────────────────────────────────────────────────────────────────────────────
# SKILL MATCHER  (TF-IDF + keyword overlap)
# ─────────────────────────────────────────────────────────────────────────────
def calculate_skill_match(candidate_skills: list, required_skills: list) -> dict:
    c_lower = [s.lower().strip() for s in candidate_skills]
    r_lower = [s.lower().strip() for s in required_skills]
    matched  = [s for s in r_lower if s in c_lower]
    missing  = [s for s in r_lower if s not in c_lower]
    pct = round(len(matched) / len(r_lower) * 100, 1) if r_lower else 0.0
    return {"matched_skills": matched, "missing_skills": missing,
            "match_pct": pct, "match_summary": f"{len(matched)}/{len(r_lower)} ({pct}%)"}


def compute_tfidf_similarity(candidate: dict, jd: dict) -> dict:
    """Cosine similarity between candidate text blob and JD skills."""
    cand_text = " ".join(candidate.get("skills", [])) + " " + candidate.get("summary", "")
    jd_text   = " ".join(jd.get("required_skills", []) + jd.get("nice_to_have", []))
    if not cand_text.strip() or not jd_text.strip():
        return {"similarity": 0.0, "similarity_pct": "0.0%"}
    vec = TfidfVectorizer().fit_transform([cand_text, jd_text])
    score = round(float(cosine_similarity(vec[0], vec[1])[0][0]) * 100, 1)
    return {"similarity": score / 100, "similarity_pct": f"{score}%"}


# ─────────────────────────────────────────────────────────────────────────────
# 5-DIMENSION SCORER  (LangChain LCEL)
# ─────────────────────────────────────────────────────────────────────────────
_SCORE_PROMPT = PromptTemplate(
    input_variables=["candidate_json", "jd_json", "skill_summary"],
    template="""You are an expert HR evaluator. Score this candidate strictly.

CANDIDATE:
{candidate_json}

JOB REQUIREMENTS:
{jd_json}

SKILL ANALYSIS:
{skill_summary}

SCORING RUBRIC (0–10):
1. Skills Match      (30%) — 0–3: <30% match | 4–6: 50–70% | 7–10: >85%
2. Experience        (25%) — 0–3: wrong domain | 4–6: adjacent | 7–10: exact fit
3. Education & Certs (15%) — 0–3: below min | 4–6: meets | 7–10: exceeds+certs
4. Projects          (20%) — 0–3: none | 4–6: 1–2 generic | 7–10: strong portfolio
5. Communication     (10%) — 0–3: poor | 4–6: adequate | 7–10: crisp+structured

Return ONLY valid JSON:
{{
  "dimension_scores": [
    {{"dimension": "Skills Match",      "score": 8, "weight": 0.30, "justification": "..."}},
    {{"dimension": "Experience",        "score": 7, "weight": 0.25, "justification": "..."}},
    {{"dimension": "Education & Certs", "score": 9, "weight": 0.15, "justification": "..."}},
    {{"dimension": "Projects",          "score": 8, "weight": 0.20, "justification": "..."}},
    {{"dimension": "Communication",     "score": 8, "weight": 0.10, "justification": "..."}}
  ],
  "strengths": ["strength 1", "strength 2"],
  "concerns":  ["concern 1"],
  "skill_gap": ["missing skill 1"]
}}

JSON:""",
)

def score_candidate(candidate: dict, jd: dict, skill_match: dict, llm) -> Optional[dict]:
    chain = _SCORE_PROMPT | llm | StrOutputParser()
    raw = chain.invoke({
        "candidate_json": json.dumps({k: candidate[k] for k in
            ("name","skills","total_experience","education","certifications","projects","summary")
            if k in candidate}, indent=2)[:1500],
        "jd_json": json.dumps({k: jd[k] for k in
            ("role_title","required_skills","min_experience","domain","seniority")
            if k in jd}, indent=2),
        "skill_summary": (
            f"Match: {skill_match['match_summary']} | "
            f"Missing: {', '.join(skill_match['missing_skills'][:5])}"
        ),
    })
    raw = raw.strip().replace("```json", "").replace("```", "").strip()
    try:
        scores = json.loads(raw)
        total = sum(d["score"] * d["weight"] for d in scores["dimension_scores"])
        scores["total_score"] = round(total, 2)
        scores["recommendation"] = (
            "HIRE"    if total >= 7.0 else
            "MAYBE"   if total >= 4.5 else
            "NO HIRE"
        )
        audit("CANDIDATE_SCORED",
              f"name={candidate.get('name')} score={scores['total_score']} rec={scores['recommendation']}")
        return scores
    except Exception as e:
        audit("SCORE_ERROR", str(e))
        return None


# ─────────────────────────────────────────────────────────────────────────────
# XAI — LIME + Dimension contributions
# ─────────────────────────────────────────────────────────────────────────────
def explain_with_lime(resume_text: str, required_skills: list, top_n: int = 6) -> list:
    """LIME explanation: which keywords most influenced the skill-match score."""
    explainer = lime.lime_text.LimeTextExplainer(class_names=["Low Fit", "High Fit"])

    def predict(texts):
        out = []
        for t in texts:
            cnt = sum(1 for s in required_skills if s.lower() in t.lower())
            p = min(cnt / max(len(required_skills), 1), 1.0)
            out.append([1 - p, p])
        return np.array(out)

    try:
        exp = explainer.explain_instance(resume_text, predict, num_features=top_n, labels=(1,))
        return exp.as_list(label=1)
    except Exception:
        return []


def explain_dimension_contributions(scores: dict) -> list:
    """Show how much each dimension pushed the score above/below the 5.0 baseline."""
    baseline = 5.0
    contribs = []
    for d in scores.get("dimension_scores", []):
        contribs.append({
            "dimension":    d["dimension"],
            "contribution": round((d["score"] - baseline) * d["weight"], 3),
        })
    return contribs


# ─────────────────────────────────────────────────────────────────────────────
# HUMAN-IN-THE-LOOP OVERRIDE
# ─────────────────────────────────────────────────────────────────────────────
VALID_DIMENSIONS = {"Skills Match", "Experience", "Education & Certs", "Projects", "Communication"}

def override_score(
    results: list,
    candidate_name: str,
    dimension: str,
    new_score: float,
    reason: str,
    reviewer: str = "HR Manager",
) -> list:
    """
    Let HR adjust any dimension score post-evaluation.
    Every change is permanently written to audit_log.json.
    """
    if dimension not in VALID_DIMENSIONS:
        print(f"❌ Invalid dimension. Choose from: {VALID_DIMENSIONS}")
        return results
    if not 0 <= new_score <= 10:
        print("❌ Score must be 0–10.")
        return results

    for r in results:
        if r["candidate_info"]["name"].lower() == candidate_name.lower():
            for d in r["scores"]["dimension_scores"]:
                if d["dimension"] == dimension:
                    original = d["score"]
                    d["score"] = new_score
                    d["justification"] += f" [OVERRIDDEN by {reviewer}: {reason}]"
                    new_total = sum(x["score"] * x["weight"] for x in r["scores"]["dimension_scores"])
                    r["scores"]["total_score"] = round(new_total, 2)
                    r["scores"]["recommendation"] = (
                        "HIRE" if new_total >= 7.0 else "MAYBE" if new_total >= 4.5 else "NO HIRE"
                    )
                    results.sort(key=lambda x: x["scores"]["total_score"], reverse=True)
                    audit("HR_OVERRIDE",
                          f"reviewer={reviewer} candidate={candidate_name} "
                          f"dim={dimension} {original}→{new_score} reason={reason}")
                    print(f"✅ Override logged | {candidate_name} | {dimension}: {original} → {new_score}")
                    return results
    print(f"❌ Candidate '{candidate_name}' not found.")
    return results


# ─────────────────────────────────────────────────────────────────────────────
# PDF REPORT GENERATOR
# ─────────────────────────────────────────────────────────────────────────────
def generate_pdf_report(results: list, jd: dict, path: str = "outputs/shortlist_report.pdf") -> str:
    """
    Professional ReportLab PDF with:
      – Cover page  – Rankings summary  – Per-candidate rubric breakdown
      – XAI (LIME keywords + dimension contributions)  – Audit footer
    """
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc = SimpleDocTemplate(
        path, pagesize=A4,
        topMargin=1.2*cm, bottomMargin=1.2*cm,
        leftMargin=2*cm, rightMargin=2*cm,
    )
    styles = getSampleStyleSheet()
    story  = []

    # ── helpers ──────────────────────────────────────────────────────────────
    def sty(name, **kw):
        return ParagraphStyle(name, **kw)

    T_TITLE   = sty("T", fontSize=22, fontName="Helvetica-Bold",
                    textColor=C_NAVY, alignment=TA_CENTER, spaceAfter=4)
    T_SUB     = sty("S", fontSize=9, textColor=C_DGRAY,
                    alignment=TA_CENTER, spaceAfter=14)
    T_SEC     = sty("SEC", fontSize=13, fontName="Helvetica-Bold",
                    textColor=C_NAVY, spaceBefore=14, spaceAfter=6)
    T_XAI     = sty("XAI", fontSize=11, fontName="Helvetica-Bold",
                    textColor=C_PURPLE, spaceBefore=10, spaceAfter=4)
    T_BODY    = sty("BD", fontSize=8.5, leading=12, textColor=C_DGRAY)
    T_FOOT    = sty("FT", fontSize=7, textColor=colors.HexColor("#9CA3AF"),
                    alignment=TA_CENTER, spaceBefore=8)

    def hr():
        return HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#D1D5DB"))

    # ── Cover ─────────────────────────────────────────────────────────────────
    story += [
        Spacer(1, 1.2*cm),
        Paragraph("TalentMatch AI", T_TITLE),
        Paragraph("HR Resume Shortlisting Report", T_TITLE),
        Spacer(1, 0.3*cm),
        hr(),
        Spacer(1, 0.3*cm),
        Paragraph(
            f"Role: <b>{jd.get('role_title','N/A')}</b> &nbsp;|&nbsp; "
            f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')} &nbsp;|&nbsp; "
            f"Model: Gemini 2.5 Flash (LangChain LCEL) &nbsp;|&nbsp; "
            f"Candidates: {len(results)}",
            T_SUB,
        ),
        Spacer(1, 0.4*cm),
        Paragraph("🔬 Powered by Explainable AI (XAI) — LIME + Dimension Contributions", T_XAI),
        Spacer(1, 0.2*cm),
        Paragraph(
            "This report was produced by an AI agent and is intended to <i>assist</i> — not replace — "
            "qualified human recruiters. Every dimension score can be overridden via <b>override_score()</b>. "
            "Full audit trail stored in <b>audit_log.json</b>.",
            T_BODY,
        ),
        PageBreak(),
    ]

    # ── Rankings summary table ────────────────────────────────────────────────
    story += [Paragraph("Candidate Rankings", T_SEC)]

    hdr = [["Rank", "Candidate", "Similarity", "Total", "Skills", "Exp", "Edu", "Projects", "Comm", "Decision"]]
    rows = hdr.copy()
    for i, r in enumerate(results, 1):
        ds = {d["dimension"]: d["score"] for d in r["scores"]["dimension_scores"]}
        rows.append([
            str(i),
            r["candidate_info"].get("name","?")[:22],
            r.get("embedding", {}).get("similarity_pct", "—"),
            f"{r['scores']['total_score']}/10",
            f"{ds.get('Skills Match',0):.1f}",
            f"{ds.get('Experience',0):.1f}",
            f"{ds.get('Education & Certs',0):.1f}",
            f"{ds.get('Projects',0):.1f}",
            f"{ds.get('Communication',0):.1f}",
            r["scores"]["recommendation"],
        ])

    tbl = Table(rows, colWidths=[0.5*cm, 3.5*cm, 1.8*cm, 1.4*cm,
                                  1.2*cm, 1.0*cm, 1.0*cm, 1.5*cm, 1.0*cm, 2.0*cm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",     (0,0), (-1,0),  C_NAVY),
        ("TEXTCOLOR",      (0,0), (-1,0),  colors.white),
        ("FONTNAME",       (0,0), (-1,0),  "Helvetica-Bold"),
        ("FONTSIZE",       (0,0), (-1,-1), 7.5),
        ("ALIGN",          (0,0), (-1,-1), "CENTER"),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [C_LGRAY, colors.white]),
        ("GRID",           (0,0), (-1,-1), 0.3, colors.HexColor("#D1D5DB")),
        ("TOPPADDING",     (0,0), (-1,-1), 4),
        ("BOTTOMPADDING",  (0,0), (-1,-1), 4),
        # top-3 highlight
        ("BACKGROUND",     (0,1), (-1,min(3,len(rows)-1)), colors.HexColor("#D1FAE5")),
    ]))
    story += [tbl, Spacer(1, 0.3*cm), PageBreak()]

    # ── Per-candidate detail ──────────────────────────────────────────────────
    for r in results:
        name  = r["candidate_info"].get("name", "?")
        total = r["scores"]["total_score"]
        rec   = r["scores"]["recommendation"]
        rank  = r.get("rank", "?")

        story += [
            Paragraph(f"#{rank}  {name}  —  {total}/10  ({rec})", T_SEC),
            hr(),
        ]

        # quick info line
        info = (
            f"Experience: {r['candidate_info'].get('total_experience','?')} yrs  |  "
            f"Skill Match: {r['skill_match']['match_summary']}  |  "
            f"Similarity: {r.get('embedding',{}).get('similarity_pct','—')}"
        )
        story.append(Paragraph(info, T_BODY))
        story.append(Spacer(1, 0.2*cm))

        # dimension table
        dim_rows = [["Dimension", "Wt", "Score", "Justification"]]
        for d in r["scores"]["dimension_scores"]:
            j = d["justification"]
            dim_rows.append([
                d["dimension"],
                f"{int(d['weight']*100)}%",
                f"{d['score']}/10",
                (j[:90] + "…") if len(j) > 90 else j,
            ])
        dt = Table(dim_rows, colWidths=[3.2*cm, 0.8*cm, 1.2*cm, 10.8*cm])
        dt.setStyle(TableStyle([
            ("BACKGROUND",     (0,0), (-1,0),  colors.HexColor("#E0F2FE")),
            ("FONTNAME",       (0,0), (-1,0),  "Helvetica-Bold"),
            ("FONTSIZE",       (0,0), (-1,-1), 7.5),
            ("ALIGN",          (1,0), (2,-1),  "CENTER"),
            ("ALIGN",          (0,0), (0,-1),  "LEFT"),
            ("ALIGN",          (3,0), (3,-1),  "LEFT"),
            ("GRID",           (0,0), (-1,-1), 0.3, colors.HexColor("#CBD5E1")),
            ("TOPPADDING",     (0,0), (-1,-1), 3),
            ("BOTTOMPADDING",  (0,0), (-1,-1), 3),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.HexColor("#FAFAFA"), colors.white]),
        ]))
        story += [dt, Spacer(1, 0.15*cm)]

        # XAI section
        story.append(Paragraph("🔬 XAI Insights", T_XAI))
        contribs = r.get("xai_contributions", [])
        if contribs:
            c_str = "  |  ".join(
                f"{c['dimension'].split()[0]}: {'+'if c['contribution']>=0 else ''}{c['contribution']:.3f}"
                for c in contribs
            )
            story.append(Paragraph(f"Dimension contributions vs baseline (5.0): <i>{c_str}</i>", T_BODY))

        lime_feats = r.get("xai_features", [])
        if lime_feats:
            l_str = "  ".join(f"'{w}'({v:+.2f})" for w, v in lime_feats[:6])
            story.append(Paragraph(f"LIME keywords: <i>{l_str}</i>", T_BODY))

        # strengths / concerns
        story.append(Spacer(1, 0.1*cm))
        if r["scores"].get("strengths"):
            story.append(Paragraph("<b>Strengths:</b> " + "  •  ".join(r["scores"]["strengths"]), T_BODY))
        if r["scores"].get("concerns"):
            story.append(Paragraph("<b>Concerns:</b> "  + "  •  ".join(r["scores"]["concerns"]),  T_BODY))
        if r["scores"].get("skill_gap"):
            story.append(Paragraph("<b>Skill Gap:</b> " + ", ".join(r["scores"]["skill_gap"]),     T_BODY))

        story += [Spacer(1, 0.4*cm), PageBreak()]

    # ── Footer disclaimer ─────────────────────────────────────────────────────
    story.append(hr())
    story.append(Paragraph(
        "AI-generated shortlist. All final hiring decisions must be made by a qualified human recruiter. "
        "Scores adjustable via override_score(). Full action trace in audit_log.json.",
        T_FOOT,
    ))

    doc.build(story)
    audit("REPORT_GENERATED", f"path={path} candidates={len(results)}")
    print(f"✅ PDF saved → {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# MAIN AGENT PIPELINE
# ─────────────────────────────────────────────────────────────────────────────
def run_agent(
    jd_text: str,
    resume_files: List[str] = None,
    linkedin_files: List[str] = None,
    llm=None,
    generate_report: bool = True,
    report_path: str = "outputs/shortlist_report.pdf",
) -> List[dict]:
    """
    Full 7-step HR Shortlisting Agent pipeline.

    Architecture: Sequential Plan-and-Execute (LangChain LCEL)
      1  Parse JD          → LangChain LCEL chain
      2  Ingest profiles   → PDF / DOCX / TXT / LinkedIn JSON
      3  Semantic match    → TF-IDF cosine + keyword overlap
      4  LLM scoring       → 5-dimension rubric via LCEL
      5  XAI               → LIME + dimension contributions
      6  Rank              → sort by total_score desc
      7  Report            → PDF + JSON

    Caching: SQLite (llm_cache.db) — zero API calls for repeated prompts.
    """
    resume_files   = resume_files   or []
    linkedin_files = linkedin_files or []

    print("\n" + "="*62)
    print("  TALENTMATCH AI  |  Gemini 2.5 Flash  |  LangChain LCEL")
    print("="*62)

    # Step 1 — Parse JD
    print("\n[1/7] Parsing Job Description…")
    jd = analyze_jd(jd_text, llm)
    print(f"   ✓ Role: {jd.get('role_title')}  |  Skills: {len(jd.get('required_skills',[]))}")

    # Step 2 — Load candidates
    print(f"\n[2/7] Loading candidates ({len(resume_files)} resumes, {len(linkedin_files)} LinkedIn)…")
    all_profiles = []
    for path in resume_files:
        try:
            text = read_resume(path)
            all_profiles.append((text, path, "resume"))
            print(f"   ✓ {os.path.basename(path)}")
        except Exception as e:
            print(f"   ✗ {path}: {e}")
    for path in linkedin_files:
        try:
            text = read_linkedin_json(path)
            all_profiles.append((text, path, "linkedin"))
            print(f"   ✓ LinkedIn: {os.path.basename(path)}")
        except Exception as e:
            print(f"   ✗ {path}: {e}")

    if not all_profiles:
        print("❌ No candidates loaded.")
        return []

    # Steps 3–5 — Per-candidate processing
    print(f"\n[3–5/7] Processing {len(all_profiles)} candidate(s)…")
    results = []
    for i, (text, path, source) in enumerate(all_profiles, 1):
        print(f"\n  ── Candidate {i}/{len(all_profiles)} ──")
        candidate = extract_candidate(text, llm)
        candidate["source"] = source
        candidate["file"]   = os.path.basename(path)

        embedding   = compute_tfidf_similarity(candidate, jd)
        skill_match = calculate_skill_match(candidate.get("skills", []), jd.get("required_skills", []))
        print(f"   Similarity: {embedding['similarity_pct']}  |  Skills: {skill_match['match_summary']}")

        scores = score_candidate(candidate, jd, skill_match, llm)
        if scores is None:
            continue

        lime_feats    = explain_with_lime(text, jd.get("required_skills", []))
        contributions = explain_dimension_contributions(scores)

        results.append({
            "candidate_info":    candidate,
            "skill_match":       skill_match,
            "embedding":         embedding,
            "scores":            scores,
            "xai_features":      lime_feats,
            "xai_contributions": contributions,
        })
        if i < len(all_profiles):
            time.sleep(DELAY_BETWEEN_CANDIDATES)

    if not results:
        print("❌ No candidates scored.")
        return []

    # Step 6 — Rank
    print("\n[6/7] Ranking…")
    results.sort(key=lambda x: x["scores"]["total_score"], reverse=True)
    for i, r in enumerate(results, 1):
        r["rank"] = i

    # Step 7 — Output
    print("\n[7/7] Generating outputs…")
    os.makedirs("outputs", exist_ok=True)
    with open("outputs/scored_candidates.json", "w") as f:
        json.dump(results, f, default=str, indent=2)
    print("   ✓ outputs/scored_candidates.json")

    if generate_report:
        generate_pdf_report(results, jd, report_path)

    # Console summary
    print("\n" + "="*62)
    print("  FINAL RANKINGS")
    print("="*62)
    print(f"  {'#':<3} {'Candidate':<24} {'Sim':<10} {'Score':<8} Decision")
    print(f"  {'-'*58}")
    for r in results:
        print(f"  #{r['rank']:<2} {r['candidate_info']['name']:<24} "
              f"{r['embedding']['similarity_pct']:<10} "
              f"{r['scores']['total_score']:<8.1f} {r['scores']['recommendation']}")

    audit("AGENT_COMPLETE", f"candidates={len(results)}")
    return results
